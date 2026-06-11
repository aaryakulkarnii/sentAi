import os
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from app.models.incident import Incident
from app.models.alert import Alert


def _ensure_dir(path: str):
    os.makedirs(os.path.dirname(path), exist_ok=True)


def generate_pdf_report(incident: Incident, alerts: list[Alert], file_path: str) -> str:
    _ensure_dir(file_path)
    
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    Story = []
    
    title_style = styles["Heading1"]
    h2_style = styles["Heading2"]
    normal_style = styles["Normal"]
    
    Story.append(Paragraph(f"Executive Incident Report: {incident.title}", title_style))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph(f"<b>Status:</b> {incident.status.upper()}", normal_style))
    Story.append(Paragraph(f"<b>Created:</b> {incident.created_at.strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    Story.append(Paragraph(f"<b>Assigned To:</b> {incident.assigned_to or 'Unassigned'}", normal_style))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("<b>Description:</b>", h2_style))
    Story.append(Paragraph(incident.description or "No description provided.", normal_style))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph(f"<b>Associated Alerts ({len(alerts)}):</b>", h2_style))
    for idx, alert in enumerate(alerts, 1):
        Story.append(Paragraph(f"{idx}. {alert.description}", styles["Heading3"]))
        Story.append(Paragraph(f"Severity: {alert.severity} | Time: {alert.created_at.strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        if alert.source_ip:
            Story.append(Paragraph(f"Source IP: {alert.source_ip}", normal_style))
        Story.append(Spacer(1, 6))
        
    doc.build(Story)
    return file_path


def generate_docx_report(incident: Incident, alerts: list[Alert], file_path: str) -> str:
    _ensure_dir(file_path)
    
    doc = Document()
    doc.add_heading(f"Executive Incident Report: {incident.title}", 0)
    
    doc.add_paragraph(f"Status: {incident.status.upper()}")
    doc.add_paragraph(f"Created: {incident.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Assigned To: {incident.assigned_to or 'Unassigned'}")
    
    doc.add_heading("Description", level=1)
    doc.add_paragraph(incident.description or "No description provided.")
    
    doc.add_heading(f"Associated Alerts ({len(alerts)})", level=1)
    
    for idx, alert in enumerate(alerts, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(alert.description).bold = True
        doc.add_paragraph(f"Severity: {alert.severity} | Time: {alert.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        if alert.source_ip:
            doc.add_paragraph(f"Source IP: {alert.source_ip}")
            
    doc.save(file_path)
    return file_path

"""Incident report generation — PDF (ReportLab) and DOCX (python-docx)."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.incident import Incident
from app.models.investigation import Investigation


async def gather_report_data(db: AsyncSession, incident_id: str) -> dict | None:
    incident = await db.get(Incident, incident_id)
    if incident is None:
        return None

    inv = (
        await db.execute(
            select(Investigation)
            .where(Investigation.incident_id == incident_id)
            .order_by(Investigation.created_at.desc())
            .limit(1)
        )
    ).scalar_one_or_none()

    timeline = (inv.attack_timeline or {}).get("events", []) if inv else []
    mitre = (inv.mitre_mappings or {}) if inv else {}
    iocs = (inv.iocs or {}).get("items", []) if inv else []
    remediation = (inv.remediation or {}) if inv else {}
    summary = remediation.get("executive_summary", "")
    steps = remediation.get("ai_recommended") or remediation.get("immediate_actions") or []

    return {
        "incident": {
            "id": incident.id,
            "title": incident.title,
            "status": incident.status,
            "risk_score": incident.risk_score,
            "created_at": incident.created_at,
        },
        "has_investigation": inv is not None,
        "executive_summary": summary,
        "timeline": timeline,
        "mitre": list(mitre.values()),
        "iocs": iocs,
        "remediation_steps": steps,
        "playbooks": remediation.get("playbooks", []),
        "generated_at": datetime.now(timezone.utc),
    }


def _risk_label(score: int) -> str:
    return "Critical" if score >= 80 else "High" if score >= 60 else "Medium" if score >= 40 else "Low"


# ── PDF ──────────────────────────────────────────────────────────────────────
def build_pdf(data: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, textColor=colors.HexColor("#111111"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, textColor=colors.HexColor("#333333"), spaceBefore=14)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10, leading=15, alignment=TA_LEFT)
    small = ParagraphStyle("small", parent=styles["BodyText"], fontSize=8, textColor=colors.grey)

    inc = data["incident"]
    elems: list = []
    elems.append(Paragraph("SentinelAI — Incident Investigation Report", h1))
    elems.append(Paragraph(f"Generated {data['generated_at'].strftime('%Y-%m-%d %H:%M UTC')}", small))
    elems.append(Spacer(1, 12))

    meta = [
        ["Incident", inc["title"]],
        ["ID", inc["id"]],
        ["Status", inc["status"].title()],
        ["Risk Score", f"{inc['risk_score']}/100  ({_risk_label(inc['risk_score'])})"],
    ]
    t = Table(meta, colWidths=[1.4 * inch, 5.1 * inch])
    t.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.grey),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
    ]))
    elems.append(t)

    elems.append(Paragraph("Executive Summary", h2))
    elems.append(Paragraph(data["executive_summary"] or "No investigation has been run for this incident.", body))

    if data["timeline"]:
        elems.append(Paragraph("Attack Timeline", h2))
        rows = [["#", "Technique", "Tactic", "Description"]]
        for i, ev in enumerate(data["timeline"], 1):
            rows.append([str(i), ev.get("technique_id") or "-", ev.get("tactic") or "-",
                         (ev.get("description") or "")[:70]])
        tt = Table(rows, colWidths=[0.3 * inch, 0.9 * inch, 1.3 * inch, 4.0 * inch])
        tt.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e0e0e0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elems.append(tt)

    if data["mitre"]:
        elems.append(Paragraph("MITRE ATT&CK Techniques", h2))
        items = [ListItem(Paragraph(f"<b>{m.get('id')}</b> — {m.get('technique')} "
                                    f"({m.get('tactic') or 'n/a'})", body)) for m in data["mitre"]]
        elems.append(ListFlowable(items, bulletType="bullet"))

    if data["iocs"]:
        elems.append(Paragraph("Indicators of Compromise", h2))
        items = [ListItem(Paragraph(f"{i.get('value')} — <i>{i.get('verdict')}</i>", body)) for i in data["iocs"]]
        elems.append(ListFlowable(items, bulletType="bullet"))

    if data["remediation_steps"]:
        elems.append(Paragraph("Recommended Remediation", h2))
        items = [ListItem(Paragraph(s, body)) for s in data["remediation_steps"]]
        elems.append(ListFlowable(items, bulletType="1"))

    elems.append(Spacer(1, 20))
    elems.append(Paragraph("Generated by SentinelAI — Autonomous SOC Platform", small))

    doc.build(elems)
    return buf.getvalue()


# ── DOCX ─────────────────────────────────────────────────────────────────────
def build_docx(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    inc = data["incident"]
    doc = Document()
    doc.add_heading("SentinelAI — Incident Investigation Report", level=0)
    p = doc.add_paragraph()
    run = p.add_run(f"Generated {data['generated_at'].strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    table = doc.add_table(rows=0, cols=2)
    for label, value in [
        ("Incident", inc["title"]),
        ("ID", inc["id"]),
        ("Status", inc["status"].title()),
        ("Risk Score", f"{inc['risk_score']}/100 ({_risk_label(inc['risk_score'])})"),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data["executive_summary"] or "No investigation has been run for this incident.")

    if data["timeline"]:
        doc.add_heading("Attack Timeline", level=1)
        for i, ev in enumerate(data["timeline"], 1):
            doc.add_paragraph(
                f"{ev.get('technique_id') or '-'} ({ev.get('tactic') or 'n/a'}): "
                f"{ev.get('description') or ''}",
                style="List Number",
            )

    if data["mitre"]:
        doc.add_heading("MITRE ATT&CK Techniques", level=1)
        for m in data["mitre"]:
            doc.add_paragraph(f"{m.get('id')} — {m.get('technique')} ({m.get('tactic') or 'n/a'})",
                              style="List Bullet")

    if data["iocs"]:
        doc.add_heading("Indicators of Compromise", level=1)
        for i in data["iocs"]:
            doc.add_paragraph(f"{i.get('value')} — {i.get('verdict')}", style="List Bullet")

    if data["remediation_steps"]:
        doc.add_heading("Recommended Remediation", level=1)
        for s in data["remediation_steps"]:
            doc.add_paragraph(s, style="List Number")

    doc.add_paragraph()
    foot = doc.add_paragraph("Generated by SentinelAI — Autonomous SOC Platform")
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(data: dict, fmt: str) -> bytes:
    if fmt == "pdf":
        return build_pdf(data)
    if fmt == "docx":
        return build_docx(data)
    raise ValueError(f"Unsupported format: {fmt}")
